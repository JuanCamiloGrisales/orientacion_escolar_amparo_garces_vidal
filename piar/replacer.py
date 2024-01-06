import os
from docx import Document
from docx2pdf import convert

def reemplazar_valores(documento, diccionario_reemplazos, new_file_path):
    """
    Reemplaza los valores especificados en un documento Word con los proporcionados en un diccionario.

    Parameters:
    - documento (str): Ruta al archivo del documento Word a procesar (plantilla).
    - diccionario_reemplazos (dict): Diccionario que contiene las sustituciones a realizar.
    - new_file_path (str): Ruta para guardar el nuevo documento con los reemplazos.

    Returns:
    new_file_path (str): Ruta del nuevo documento con los reemplazos en pdf.
    """
    index = 0

    # Abre el documento Word
    documento = Document(documento)

    def get_text_from_element(element):
        text = ""
        if element.tag.endswith('tbl'):
            # Procesa tablas
            for row in element:
                for cell in row:
                    for paragraph in cell:
                        for run in paragraph:
                            if run.text:
                                text += run.text
        else:
            # Procesa parrafos
            for paragraph in element:   
                for run in paragraph:
                    if run.text:
                        text += run.text
        return text

    for antigua, nueva in diccionario_reemplazos.items():
        if not isinstance(nueva, str):
            for element in documento.element.body:
                index += 1
                text = get_text_from_element(element)
                if antigua in text:
                    doc2 = nueva
                    for element in doc2.element.body:
                        documento.element.body.insert(index, element)
                        index += 1
            
            for p in documento.paragraphs:
                if antigua in p.text:
                    p.text = p.text.replace(antigua, '')
            for table in documento.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if antigua in cell.text:
                            cell.text = cell.text.replace(antigua, '')


    # Itera sobre los párrafos del documento
    for p in documento.paragraphs:
        # Itera sobre las sustituciones a realizar
        for antigua, nueva in diccionario_reemplazos.items():
            # Verifica si la cadena antigua está presente en el texto del párrafo
            if antigua in p.text and isinstance(nueva, str):
                # Realiza el reemplazo en el texto del párrafo
                p.text = p.text.replace(antigua, nueva)

    # Itera sobre las tablas, filas y celdas del documento
    for table in documento.tables:
        for row in table.rows:
            for cell in row.cells:
                # Itera sobre las sustituciones a realizar
                for antigua, nueva in diccionario_reemplazos.items():
                    # Verifica si la cadena antigua está presente en el texto de la celda
                    if antigua in cell.text and isinstance(nueva, str):
                        # Realiza el reemplazo en el texto de la celda
                        cell.text = cell.text.replace(antigua, nueva)

    # Save the document and convert it
    documento.save(new_file_path)
    # convert(new_file_path)
    # os.remove(new_file_path)
    # return new_file_path.replace(".docx", ".pdf")
    return new_file_path